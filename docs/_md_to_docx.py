"""Convert service-plan.md → service-plan.docx (한국어 친화 폰트/스타일)."""
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOCS_DIR = Path(__file__).parent

# 파일별 표지 (제목, 부제)
TITLES = {
    "service-plan.md":     ("Lon — AI 사업제안서 자동 생성기", "서비스 기획서"),
    "screen-design.md":    ("Lon — AI 사업제안서 자동 생성기", "화면 설계서"),
    "db-design.md":        ("Lon — AI 사업제안서 자동 생성기", "DB 설계서 (Hybrid: MariaDB + MongoDB)"),
    "interface-spec.md":   ("Lon — AI 사업제안서 자동 생성기", "연계 정의서"),
    "menu-structure.md":   ("Lon — AI 사업제안서 자동 생성기", "메뉴 구성도"),
    "dev-setup-guide.md":  ("Lon — AI 사업제안서 자동 생성기", "개발환경 구성 가이드"),
    "운영매뉴얼.md":         ("Lon — AI 사업제안서 자동 생성기", "운영 매뉴얼 (개발환경/사용법/기동테스트)"),
}

KO_FONT = "맑은 고딕"
EN_FONT = "Calibri"
MONO_FONT = "Consolas"


def set_run_font(run, size=None, bold=None, color=None, mono=False):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    font = MONO_FONT if mono else EN_FONT
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), MONO_FONT if mono else KO_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline(paragraph, text):
    """Render inline **bold** and `code` segments."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, mono=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    sizes = {1: 20, 2: 16, 3: 13, 4: 12}
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=True,
                 color=RGBColor(0x1F, 0x3A, 0x5F) if level <= 2 else None)
    return p


def add_paragraph(doc, text, bullet=False, ordered_index=None):
    p = doc.add_paragraph()
    if bullet:
        p.paragraph_format.left_indent = Cm(0.6)
        prefix_run = p.add_run("• ")
        set_run_font(prefix_run)
    elif ordered_index is not None:
        p.paragraph_format.left_indent = Cm(0.6)
        prefix_run = p.add_run(f"{ordered_index}. ")
        set_run_font(prefix_run)
    add_inline(p, text)
    return p


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # 음영
    pPr = p._p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    text = "\n".join(lines)
    run = p.add_run(text)
    set_run_font(run, size=9, mono=True)


def parse_table(rows):
    """rows: list of '| a | b |' lines (헤더 + 구분 + 바디)."""
    def split_row(line):
        line = line.strip()
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]
        return [c.strip() for c in line.split("|")]

    header = split_row(rows[0])
    # rows[1] is separator
    body = [split_row(r) for r in rows[2:]]
    return header, body


def add_table(doc, header, body):
    cols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    # Header
    for i, txt in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # 헤더 배경색
        from docx.oxml import OxmlElement
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3A5F")
        tcPr.append(shd)
    # Body
    for r, row in enumerate(body, start=1):
        for c, txt in enumerate(row):
            if c >= cols:
                continue
            cell = table.rows[r].cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            add_inline(p, txt)
            for run in p.runs:
                set_run_font(run, size=10)
    doc.add_paragraph()


def convert(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 코드 블록
        if stripped.startswith("```"):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            add_code_block(doc, buf)
            i += 1
            continue

        # 표
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|\s*$", lines[i + 1].strip()):
            tbl_lines = [stripped]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            header, body = parse_table(tbl_lines)
            add_table(doc, header, body)
            continue

        # 인용 (> )
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            run_q = p.add_run("│ ")
            set_run_font(run_q, color=RGBColor(0x1F, 0x3A, 0x5F), bold=True)
            add_inline(p, stripped[2:])
            for run in p.runs[1:]:
                run.italic = True
            i += 1
            continue

        # 헤딩
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            add_heading(doc, m.group(2), len(m.group(1)))
            i += 1
            continue

        # 수평선
        if stripped == "---":
            p = doc.add_paragraph()
            run = p.add_run("─" * 40)
            set_run_font(run, color=RGBColor(0xBB, 0xBB, 0xBB))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # 불릿
        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            add_paragraph(doc, text, bullet=True)
            i += 1
            continue

        # 순서 목록
        m = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if m:
            add_paragraph(doc, m.group(2), ordered_index=m.group(1))
            i += 1
            continue

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # 일반 문단
        add_paragraph(doc, stripped)
        i += 1


def build_doc(md_path: Path, title: str, subtitle: str) -> Document:
    doc = Document()
    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    rfonts.set(qn("w:eastAsia"), KO_FONT)

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # 표지
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(60)
    run = title_p.add_run(title)
    set_run_font(run, size=24, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(subtitle)
    set_run_font(run, size=18, bold=True, color=RGBColor(0x33, 0x33, 0x33))

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_before = Pt(24)
    run = meta_p.add_run("v0.1 (초안)   |   2026-05-03   |   파일: " + md_path.name)
    set_run_font(run, size=11, color=RGBColor(0x88, 0x88, 0x88))

    doc.add_page_break()

    convert(md_path.read_text(encoding="utf-8"), doc)
    return doc


def main():
    targets = [(name, *TITLES[name]) for name in TITLES if (DOCS_DIR / name).exists()]
    if not targets:
        print("No markdown files found.")
        return
    for name, title, subtitle in targets:
        src = DOCS_DIR / name
        dst = src.with_suffix(".docx")
        doc = build_doc(src, title, subtitle)
        doc.save(dst)
        print(f"OK  {name}  ->  {dst.name}")


if __name__ == "__main__":
    main()
